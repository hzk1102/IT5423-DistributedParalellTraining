"""Build a detailed, beginner-friendly Word report (.docx) for Topic 2.

    uv run --no-project python bench/make_docx.py

Reads the measured data in results/ and the figures in results/figures/, and
writes results/Pipeline_Parallelism_Report.docx. The prose is written so a
reader with no machine-learning background can follow it (analogies + a glossary)
while still containing the full technical detail and every measured number.

If you change numbers, edit REPORT.md (the source of truth) and mirror them here.
"""
from __future__ import annotations

import csv
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "results", "figures")
RESULTS_CSV = os.path.join(ROOT, "results", "results.csv")
OUT = os.path.join(ROOT, "results", "Pipeline_Parallelism_Report.docx")

ACCENT = RGBColor(0x1F, 0x47, 0x88)  # dark blue for headings


# --------------------------------------------------------------------------- #
# small helpers                                                               #
# --------------------------------------------------------------------------- #
def body(doc, text, size=11, italic=False, bold=False, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):  # (lead, rest) -> bold lead
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(12)


def figure(doc, filename, cap, width=6.1):
    path = os.path.join(FIG, filename)
    if not os.path.exists(path):
        body(doc, f"[figure {filename} not found — run bench/plots.py and bench/plot_analysis.py]",
             italic=True)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        _shade(hdr[i], "1F4788")
        for rr in hdr[i].paragraphs[0].runs:
            rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click here and choose “Update Field” to build the contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)


def quote_box(doc, title, text):
    """A shaded 'note' paragraph for honesty boxes / important caveats."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shade(cell, "FBF3D6")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    r.bold = True
    p.add_run(text)
    for rr in p.runs:
        rr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# --------------------------------------------------------------------------- #
# build the document                                                          #
# --------------------------------------------------------------------------- #
def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- title page ---- #
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Pipeline Parallelism for Models\nExceeding Single-GPU Memory")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("A hands-on benchmark study: DeepSpeed Pipeline Parallelism vs. "
                     "PyTorch torch.distributed.pipelining vs. single-GPU gradient checkpointing, "
                     "on two NVIDIA T4 GPUs")
    rs.italic = True; rs.font.size = Pt(13)

    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Course: Distributed & Parallel Training (GenAI Master) — Topic 2\n").bold = True
    meta.add_run("Model: gpt2-xl (1.5 billion parameters)   •   Dataset: WikiText\n")
    meta.add_run("Team members: ____________________________\n")
    meta.add_run("Date: 26 June 2026")
    doc.add_page_break()

    # ---- contents ---- #
    heading(doc, "Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ====================================================================== #
    heading(doc, "1. Executive Summary", level=1)
    body(doc, "Modern language models are often too large to fit inside a single graphics card "
              "(GPU). This project studies one of the standard solutions — pipeline parallelism — "
              "which splits a model across several GPUs like stations on a factory assembly line, "
              "so that a model too big for one GPU can still be trained.")
    body(doc, "We trained the 1.5-billion-parameter model gpt2-xl on two NVIDIA T4 GPUs (16 GB "
              "each) using two different software frameworks for pipeline parallelism, and compared "
              "them against the alternative of squeezing the model onto a single GPU using "
              "memory-saving tricks. We measured speed (tokens processed per second), memory use, "
              "the pipeline's idle time (the 'bubble'), and training quality. In plain terms, here "
              "is what we found:")
    bullets(doc, [
        ("It works. ", "A model that does not fit on one 16 GB GPU (opt-2.7b) trains successfully "
         "across two GPUs with the pipeline, using about 10.6 GB on each GPU."),
        ("More micro-batches = less idle time. ", "Feeding the pipeline in more, smaller pieces "
         "shrinks the idle 'bubble' almost exactly as theory predicts, and speed rises from "
         "1,281 to 1,959 tokens/second."),
        ("Newer is not automatically faster. ", "The newest 'zero-bubble' scheduling trick gives "
         "only about a 6% speed-up on our hardware — because our two GPUs are connected by a "
         "relatively slow link (PCIe, no NVLink) that cancels most of the benefit."),
        ("PyTorch beat DeepSpeed here. ", "PyTorch's pipeline was faster (1,835 vs. 1,656 "
         "tokens/second) and used less memory (8.4 vs. 11.5 GB per GPU); DeepSpeed repays this "
         "with more automation and convenience."),
        ("Two GPUs vs one. ", "The two-GPU pipeline was about 1.5× faster than one GPU and used "
         "less memory per GPU — and, most importantly, it can train models a single GPU cannot "
         "hold at all."),
    ])
    body(doc, "We also report, openly, a measurement bug we found in the single-GPU baseline and "
              "exactly how we corrected for it (Section 9). The headline conclusion: on this "
              "hardware, pipeline parallelism is primarily a way to fit bigger models (a memory "
              "win) and a modest speed win — not a magic 2× speed-up.")

    # ====================================================================== #
    heading(doc, "2. Plain-Language Primer (Key Terms)", level=1)
    body(doc, "If you are not familiar with machine learning, read this section first. Every term "
              "used later is explained here with an everyday analogy.")
    glossary = [
        ("GPU (graphics card): ", "a very fast processor good at the heavy arithmetic that "
         "training uses. We use two NVIDIA T4 GPUs."),
        ("VRAM (GPU memory): ", "the GPU's working memory. Ours is 16 GB per GPU. If a job needs "
         "more than this, it crashes with an 'out of memory' (OOM) error."),
        ("Model & parameters: ", "the model is the network we train; its 'parameters' are the "
         "billions of numbers it adjusts to learn. gpt2-xl has 1.5 billion of them."),
        ("Training / fine-tuning: ", "showing the model text and nudging its parameters so it "
         "predicts the next word better. Fine-tuning is training a model that already exists."),
        ("Optimizer (Adam): ", "the bookkeeper that decides how to nudge each parameter. Adam "
         "keeps two extra numbers per parameter, which is why training needs far more memory than "
         "just storing the model."),
        ("fp16 (half precision): ", "storing numbers with 16 bits instead of 32 to save memory and "
         "run faster. Our T4 GPUs support fp16 but not the newer 'bf16' format."),
        ("Throughput (tokens/second): ", "our main speed measure — how many words-pieces "
         "('tokens') the system processes per second. Higher is better."),
        ("Pipeline parallelism: ", "splitting the model into consecutive chunks ('stages'), one "
         "per GPU, and streaming the data through them like an assembly line (see Section 4.3)."),
        ("The bubble: ", "the idle time at the start and end of the pipeline when some GPUs have "
         "no work yet — wasted capacity we try to minimise (Section 4.4)."),
        ("Micro-batch: ", "one small piece of a training batch. Splitting a batch into more "
         "micro-batches keeps the assembly line fuller and shrinks the bubble."),
        ("Gradient checkpointing: ", "a memory-saving trick for a single GPU: don't store all "
         "intermediate results, recompute them later. Saves memory, costs extra computation."),
        ("DeepSpeed / PyTorch pipelining: ", "the two software libraries we compare for doing "
         "pipeline parallelism."),
        ("Perplexity: ", "a quality score for language models — lower means the model is less "
         "'surprised' by real text, i.e. better."),
        ("MFU (Model FLOPs Utilization): ", "what fraction of the GPU's theoretical peak speed we "
         "actually achieve. A normalised efficiency score; higher is better."),
    ]
    bullets(doc, glossary)

    # ====================================================================== #
    heading(doc, "3. Introduction", level=1)
    heading(doc, "3.1 The problem", level=2)
    body(doc, "Training a large language model with the Adam optimizer costs roughly 16 bytes of "
              "GPU memory per parameter: 2 bytes for the model weights, 2 for their gradients, "
              "4 for a high-precision master copy, and 4 + 4 for Adam's two bookkeeping numbers. "
              "For gpt2-xl's 1.5 billion parameters that is about 24 GB — already more than a "
              "single 16 GB T4 can hold, before we even store the intermediate activations. So we "
              "must either shrink the memory footprint to fit one GPU, or split the model across "
              "GPUs. This project measures both and compares them.")
    heading(doc, "3.2 What the assignment asks", level=2)
    body(doc, "Topic 2 sets four concrete questions, which structure our results (Section 7):")
    bullets(doc, [
        ("Deliverable 1 — Enablement: ", "fine-tune a model that does not fit on one GPU, using "
         "pipeline parallelism on two GPUs."),
        ("Deliverable 2 — Framework comparison: ", "compare DeepSpeed Pipeline Parallelism against "
         "the PyTorch-native pipeline."),
        ("Deliverable 3 — The bubble: ", "analyse the pipeline 'bubble' and how the number of "
         "micro-batches affects performance."),
        ("Deliverable 4 — Two GPUs vs one: ", "compare the two-GPU pipeline against a single GPU "
         "using gradient checkpointing (speed and memory)."),
    ])
    body(doc, "Note on the tools: the assignment brief names an old PyTorch API "
              "(torch.distributed.pipeline). That API is now deprecated; we use its modern "
              "replacement, torch.distributed.pipelining. Documenting this is itself part of "
              "Deliverable 2.")
    heading(doc, "3.3 Our hardware, and why it shapes every result", level=2)
    body(doc, "We ran on Kaggle's free 2× T4 environment. Two facts about it drive our findings:")
    bullets(doc, [
        ("The two GPUs talk over PCIe with no NVLink. ", "This is a relatively slow connection. "
         "Pipeline parallelism constantly sends data between the GPUs, so this slow link is a real, "
         "measurable cost — and it is why the newest scheduling tricks help less than on "
         "expensive datacentre hardware."),
        ("The T4 is an older 'Turing' GPU. ", "It supports fp16 but not bf16, and not the fastest "
         "attention kernels. So we use fp16 with 'loss scaling' for numerical stability."),
        ("Only two pipeline stages (P = 2). ", "With just two GPUs, the best possible speed-up is "
         "limited, and the idle bubble is large unless we use many micro-batches. This is exactly "
         "the regime the assignment asks us to probe."),
    ])

    # ====================================================================== #
    heading(doc, "4. Background: Training Big Models Across GPUs", level=1)
    heading(doc, "4.1 Why the model doesn't fit (the memory math)", level=2)
    body(doc, "The table below shows why gpt2-xl needs special handling and why the much larger "
              "opt-6.7b is out of reach even when split. 'Full-Adam' assumes the 16 bytes/parameter "
              "above.")
    table(doc,
          ["Model", "Parameters", "Full-Adam memory", "Fits one 16 GB T4?"],
          [["distilgpt2", "82 million", "~1.3 GB", "Yes (too small to be interesting)"],
           ["gpt2-xl (our model)", "1.5 billion", "~24 GB", "No — needs splitting or tricks"],
           ["opt-2.7b (enablement)", "2.7 billion", "~43 GB", "No — even with memory tricks"],
           ["opt-6.7b", "6.7 billion", "~107 GB", "No — infeasible even on two T4s"]])

    heading(doc, "4.2 Ways to split work across GPUs", level=2)
    bullets(doc, [
        ("Data parallelism: ", "copy the whole model onto each GPU and split the data. Does not "
         "help when the model itself is too big for one GPU."),
        ("Tensor parallelism: ", "split individual layers across GPUs. Very communication-heavy; "
         "usually needs a fast NVLink connection."),
        ("Pipeline parallelism (this project): ", "split the model by depth into stages, one per "
         "GPU. Communication happens only at the stage boundaries — a good fit for a slow link."),
    ])

    heading(doc, "4.3 Pipeline parallelism — the assembly line", level=2)
    body(doc, "Imagine building a car on an assembly line with two stations. Station 1 installs the "
              "engine, then passes the car to Station 2, which installs the wheels. If only one car "
              "ever moves down the line, Station 2 sits idle while Station 1 works, and vice-versa — "
              "half the factory is always idle. But if a steady stream of cars flows through, both "
              "stations stay busy at the same time.")
    body(doc, "Pipeline parallelism does exactly this with a neural network. GPU 1 holds the first "
              "half of the model's layers (Station 1); GPU 2 holds the second half (Station 2). A "
              "batch of data is split into 'micro-batches' (the cars). As GPU 2 processes "
              "micro-batch 1, GPU 1 is already working on micro-batch 2, and so on. The more "
              "micro-batches we stream through, the busier both GPUs stay.")

    heading(doc, "4.4 The 'bubble' — unavoidable idle time", level=2)
    body(doc, "At the very start, GPU 2 must wait for the first micro-batch to come out of GPU 1; "
              "at the very end, GPU 1 has finished but GPU 2 is still draining the last "
              "micro-batches. This start-up and drain idle time is called the bubble. For a "
              "pipeline with P stages and M micro-batches, the fraction of time wasted is:")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run("bubble fraction  =  (P − 1) / (M + P − 1)")
    rr.bold = True; rr.font.size = Pt(13)
    body(doc, "With our two GPUs (P = 2) this becomes 1 / (M + 1). The more micro-batches M we use, "
              "the smaller the bubble:")
    table(doc, ["Micro-batches M", "1", "2", "4", "8", "16"],
          [["Idle bubble", "50%", "33%", "20%", "11%", "5.9%"]])
    body(doc, "So we expect speed to rise as we add micro-batches — which is exactly what we test "
              "in Section 7.2. The catch: more micro-batches can use more memory and add "
              "communication overhead, so the gains taper off.")

    heading(doc, "4.5 The schedule ladder (how the bubble was attacked over time)", level=2)
    body(doc, "Researchers have invented progressively cleverer 'schedules' for ordering the "
              "forward and backward work to shrink the bubble. We test this whole ladder:")
    table(doc,
          ["Schedule", "Idea", "Effect on bubble"],
          [["GPipe (2019)", "all forward passes, then all backward passes", "baseline; uses lots of memory"],
           ["1F1B", "alternate one-forward-one-backward", "same bubble as GPipe, but far less memory"],
           ["Interleaved 1F1B", "give each GPU several non-adjacent chunks", "smaller bubble, more communication"],
           ["Zero-Bubble (2024)", "split the backward pass to fill the gaps", "bubble ≈ 0 (in theory)"],
           ["DualPipe (2025)", "two-directional full overlap", "frontier; for huge clusters only"]])
    body(doc, "We implement and measure up to Zero-Bubble. DualPipe is the current state of the art "
              "but targets massive multi-node clusters, so we cite it as related work rather than "
              "running it.")

    heading(doc, "4.6 The single-GPU survival kit (our baseline's tools)", level=2)
    bullets(doc, [
        ("Gradient checkpointing: ", "instead of storing every intermediate result during the "
         "forward pass, store only a few and recompute the rest during the backward pass. Trades "
         "about 30% extra computation for a large memory saving — the price of fitting on one GPU."),
        ("8-bit Adam: ", "store the optimizer's bookkeeping numbers in 8 bits instead of 32, "
         "cutting their memory by ~4×. We use this everywhere so the comparison is fair."),
    ])

    # ====================================================================== #
    heading(doc, "5. Experimental Setup", level=1)
    heading(doc, "5.1 Hardware & software", level=2)
    bullets(doc, [
        ("GPUs: ", "2× NVIDIA T4 (16 GB), Kaggle, connected by PCIe (no NVLink)."),
        ("Precision: ", "fp16 throughout, with the loss computed in higher precision for stability."),
        ("PyTorch: ", "version 2.10.0+cu128 (recorded in every result row) — new enough to include "
         "the zero-bubble schedules."),
    ])
    heading(doc, "5.2 Models", level=2)
    body(doc, "gpt2-xl (1.5B) is our primary model: too big for one GPU under full Adam, but it "
              "splits neatly across two. opt-2.7b is the 'enablement' demonstration (it does not "
              "fit on one GPU at all). Small models (gpt2, distilgpt2) are used only to verify "
              "correctness.")
    heading(doc, "5.3 Dataset", level=2)
    body(doc, "WikiText (the standard Salesforce/wikitext benchmark), tokenised with each model's "
              "own tokenizer and packed into fixed blocks of 512 tokens. The task is next-word "
              "prediction. The same data, in the same order, feeds all three systems so the "
              "comparison is fair.")
    heading(doc, "5.4 The three systems we compare", level=2)
    table(doc,
          ["", "Single GPU (baseline)", "PyTorch pipelining", "DeepSpeed PP"],
          [["GPUs", "1", "2", "2"],
           ["How it fits", "checkpointing + 8-bit Adam", "split model (2 stages)", "split model (2 stages)"],
           ["Schedules", "n/a", "GPipe → 1F1B → Interleaved → Zero-Bubble", "1F1B (built in)"],
           ["Loss scaling", "manual", "manual", "automatic (dynamic)"],
           ["Perplexity eval", "yes", "not implemented", "yes"]])
    heading(doc, "5.5 What we measure", level=2)
    bullets(doc, [
        ("Throughput (tokens/sec): ", "main speed metric, measured after dropping the first few "
         "warm-up steps."),
        ("Peak memory per GPU (GB): ", "the most any single GPU used — the number that decides "
         "whether a job fits."),
        ("Bubble: ", "both the theoretical formula and a measured estimate from the throughput."),
        ("MFU: ", "achieved fraction of the T4's peak speed."),
        ("Convergence: ", "training loss over time, and perplexity, to confirm the model actually "
         "learns."),
    ])
    heading(doc, "5.6 The experiment list", level=2)
    bullets(doc, [
        "Enablement: opt-2.7b on one GPU (expected to fail) vs. two GPUs (expected to succeed).",
        "Micro-batch sweep: M = 1, 2, 4, 8, 16 for both pipeline frameworks.",
        "Schedule ladder: GPipe, 1F1B, Interleaved 1F1B, and two Zero-Bubble variants at M = 8.",
        "Two GPUs vs one GPU at the matched setting (M = 8).",
        "Convergence: a longer 300-step run for each system.",
    ])

    # ====================================================================== #
    heading(doc, "6. How Our System Works (Methodology)", level=1)
    heading(doc, "6.1 Splitting the model into stages", level=2)
    body(doc, "A Hugging Face model is one big object; the pipeline frameworks need it as an "
              "ordered list of pieces that each take one tensor in and pass one tensor out. We "
              "wrote one shared module (model_split.py) that turns gpt2-xl into: [word + position "
              "embeddings] → [transformer block 1] → … → [block 48] → [final layer + output head]. "
              "Both frameworks and the single-GPU baseline use this same decomposition, so they "
              "train the same model.")
    heading(doc, "6.2 A hidden trap: the causal mask", level=2)
    body(doc, "A language model must only look at past words, never future ones, when predicting "
              "the next word. In recent versions of the Transformers library this 'don't peek "
              "ahead' rule is applied by the top-level model, not by each block. When we call the "
              "blocks individually (as the pipeline does), that rule would be silently lost — the "
              "model could cheat by looking at future words, ruining training without any error "
              "message. We fix this by explicitly building and passing the 'no peeking' mask into "
              "every block, and we verify the whole split reproduces the original model's outputs "
              "(check_split.py) before trusting any result.")
    heading(doc, "6.3 Keeping the comparison fair", level=2)
    body(doc, "All three systems share the same dataset, tokenizer, sequence length (512), random "
              "seed, fp16 setting, optimizer, and number of training steps. The only thing that "
              "differs is how the model is placed on the GPU(s) — which is exactly what we want to "
              "measure.")

    # ====================================================================== #
    heading(doc, "7. Results", level=1)
    body(doc, "All numbers below come from the measured run logged in results/results.csv and "
              "results/loss_curves.csv on Kaggle 2× T4.")

    heading(doc, "7.1 Deliverable 1 — Enablement: fitting a model one GPU can't hold", level=2)
    table(doc,
          ["Model", "One T4 (with memory tricks)", "Two T4 (pipeline)"],
          [["opt-2.7b (2.7B)", "OUT OF MEMORY — does not fit",
            "Trains: 1,701 tokens/s, 10.6 GB/GPU"]])
    body(doc, "This is the headline result: the pipeline holds half the model on each GPU "
              "(~10.6 GB each) and trains opt-2.7b successfully, whereas a single 16 GB T4 runs out "
              "of memory before training can start. (The single-GPU attempt crashes, so it leaves "
              "no row in the results file — the crash itself is the evidence, and should be "
              "screenshotted from the notebook log for the appendix.)")

    heading(doc, "7.2 Deliverable 3 — Micro-batches and the bubble", level=2)
    body(doc, "As predicted in Section 4.4, adding micro-batches shrinks the bubble and raises "
              "throughput. The table shows the PyTorch 1F1B pipeline and DeepSpeed; note also how "
              "GPipe's memory grows with M until it runs out of memory at M = 8.")
    table(doc,
          ["M", "Theory bubble", "PyTorch tok/s", "PyTorch GB/GPU", "DeepSpeed tok/s", "GPipe GB/GPU"],
          [["1", "50%", "—", "—", "689", "—"],
           ["2", "33%", "1,281", "6.9", "1,015", "7.3"],
           ["4", "20%", "1,605", "8.4", "1,353", "11.1"],
           ["8", "11%", "1,835", "8.4", "1,656", "OUT OF MEMORY"],
           ["16", "5.9%", "1,959", "8.4", "1,869", "n/a"]])
    figure(doc, "throughput_vs_M.png",
           "Figure 1. Throughput rises with the number of micro-batches and begins to level off, "
           "for both frameworks — the shape the bubble formula predicts. PyTorch leads at every M.")
    figure(doc, "bubble_empirical.png",
           "Figure 2. The measured bubble (derived from throughput) sits slightly above the "
           "theoretical curve. The small gap is the real cost of sending data over the slow PCIe "
           "link and launching each micro-batch — overheads the idealised formula ignores.")
    figure(doc, "memory_vs_M.png",
           "Figure 3. The key memory result: 1F1B keeps memory flat (~8.4 GB) as M grows, while "
           "GPipe's memory climbs with M and exceeds the 16 GB limit at M = 8. Same bubble, very "
           "different memory.")

    heading(doc, "7.3 The schedule ladder — does 'newest' mean 'fastest'?", level=2)
    body(doc, "We compared the whole modern ladder at M = 8. The honest answer for our hardware: "
              "barely. Interleaving gives essentially no speed-up (its extra communication cancels "
              "the smaller bubble), and only the most advanced 'ZBV' zero-bubble schedule gives a "
              "meaningful ~6% — far short of the theoretical 'bubble ≈ 0', because our slow PCIe "
              "link cannot hide the extra work.")
    table(doc,
          ["Schedule", "Throughput (tok/s)", "Memory/GPU (GB)", "vs. 1F1B"],
          [["GPipe", "OUT OF MEMORY at M=8", "—", "—"],
           ["1F1B", "1,835", "8.4", "baseline"],
           ["Interleaved 1F1B", "1,836", "9.3", "+0.1%"],
           ["Interleaved Zero-Bubble", "1,856", "9.5", "+1.1%"],
           ["ZBV (Zero-Bubble)", "1,949", "9.3", "+6.2%"]])
    figure(doc, "schedule_ladder.png",
           "Figure 4. Best throughput per schedule (PyTorch). The newest zero-bubble schedule is "
           "fastest, but the gains over plain 1F1B are modest on two PCIe-connected T4s.")

    heading(doc, "7.4 Deliverable 2 — DeepSpeed vs. PyTorch", level=2)
    body(doc, "At the matched setting (M = 8), PyTorch's pipeline was faster and leaner; DeepSpeed "
              "trades that for convenience features.")
    table(doc,
          ["Property", "PyTorch pipelining", "DeepSpeed PP"],
          [["Throughput (tok/s)", "1,835", "1,656"],
           ["Memory/GPU (GB)", "8.4", "11.5"],
           ["Efficiency (MFU)", "14.5%", "13.1%"],
           ["Final loss (300 steps)", "2.65", "2.76"],
           ["Schedules", "full ladder (GPipe→ZBV)", "1F1B only"],
           ["Loss scaling", "manual", "automatic"],
           ["Perplexity eval", "not implemented", "yes (14.99)"]])
    figure(doc, "system_comparison.png",
           "Figure 5. Best throughput (left) and peak memory per GPU (right) by system. The red "
           "dotted line marks the 16 GB limit of one T4.")
    body(doc, "A practical note: on the DeepSpeed run we observed a warning that the inter-GPU "
              "send/receive operations were being serialised — a direct symptom of the slow PCIe "
              "link, and itself a small piece of evidence for why communication overhead matters "
              "in this study.")

    heading(doc, "7.5 Deliverable 4 — Two GPUs vs. one GPU", level=2)
    quote_box(doc, "Honesty note (measurement correction).",
              "Our single-GPU baseline had a data-loading bug: it accidentally trained on 8× the "
              "intended amount of data per step, and its speedometer under-counted by the same 8×. "
              "Because the elapsed time and the work actually done are both known, the true speed "
              "is recovered exactly by multiplying by 8 (152 → ~1,221 tokens/s). The remaining "
              "imperfection — its data batches were larger than the pipeline's — actually flatters "
              "the single GPU, so the pipeline's lead below is a conservative (lower-bound) "
              "estimate. The bug is fixed in the code for any future re-run.")
    table(doc,
          ["System", "Throughput (tok/s)", "Memory/GPU (GB)"],
          [["One GPU, no checkpointing", "OUT OF MEMORY", "—"],
           ["One GPU, + checkpointing", "~1,221 (corrected)", "12.2"],
           ["Two GPUs, PyTorch 1F1B", "1,835", "8.4"],
           ["Two GPUs, DeepSpeed", "1,656", "11.5"]])
    body(doc, "Reading this: the two-GPU pipeline is about 1.5× faster than the corrected "
              "single-GPU baseline, and uses less memory per GPU because each GPU holds only half "
              "the model. The speed-up is real but well below a perfect 2× — the bubble and the "
              "slow link eat into it. The pipeline's biggest advantage is not raw speed but the "
              "ability to train models that simply do not fit on one GPU (Section 7.1).")

    heading(doc, "7.6 Convergence — does it actually learn?", level=2)
    body(doc, "Yes. All systems reduce the training loss smoothly. The fair comparison is between "
              "the two pipeline frameworks, which used identical batches: PyTorch reaches a final "
              "loss of 2.65 and DeepSpeed 2.76 — within a hair of each other, confirming they solve "
              "the same problem and differ only in cost, not correctness.")
    figure(doc, "convergence.png",
           "Figure 6. Training loss over optimizer steps (bold = smoothed trend, faint = raw). "
           "The single-GPU curve looks lower, but that is misleading — see Figure 7.")
    body(doc, "Figure 6 makes the single GPU look best, but that is the 8×-data bug again: per "
              "step it saw 8× more text. Figure 7 fixes this by plotting against the number of "
              "tokens actually processed. On that fair axis, at an equal amount of data (~1.2 "
              "million tokens, where the pipeline runs end), the single GPU is at about 2.72 — "
              "actually slightly worse than PyTorch's 2.62. In other words, per unit of data the "
              "pipeline learns just as efficiently.")
    figure(doc, "convergence_vs_tokens.png",
           "Figure 7. The fair view: training loss vs. tokens actually processed. The single GPU "
           "(orange) only reaches a lower loss because it consumed far more data; compared at the "
           "same amount of data, the pipeline is at least as good.")

    # ====================================================================== #
    heading(doc, "8. Discussion — What It All Means", level=1)
    bullets(doc, [
        ("The bubble theory is confirmed. ", "Throughput follows the (P−1)/(M+P−1) curve, with a "
         "small, explainable extra cost from the PCIe link. This is a textbook result reproduced "
         "on real, cheap hardware."),
        ("1F1B vs GPipe is about memory, not speed. ", "They have the same bubble, but 1F1B keeps "
         "memory flat while GPipe overflows — so 1F1B is simply the better default."),
        ("Cutting-edge schedules are interconnect-bound. ", "Zero-bubble's promise depends on a "
         "fast link to hide its extra work. On PCIe T4s it gives only ~6%. A valuable, honest "
         "lesson: a technique's headline benefit can evaporate on commodity hardware."),
        ("Framework choice is a real trade-off. ", "PyTorch is faster and leaner and exposes more "
         "schedules; DeepSpeed is more automated (loss scaling, partitioning, evaluation). For a "
         "research study PyTorch was more flexible; for a quick robust run DeepSpeed is easier."),
        ("Pipeline parallelism = mostly a memory tool here. ", "At two GPUs over PCIe it delivers "
         "a ~1.5× speed-up and a per-GPU memory reduction, but its decisive value is enabling "
         "models that one GPU cannot hold at all."),
    ])

    # ====================================================================== #
    heading(doc, "9. Honesty & Limitations", level=1)
    body(doc, "A good benchmark study is honest about what it cannot claim. The main caveats:")
    bullets(doc, [
        ("Single-GPU batching bug (corrected, see Section 7.5). ", "Its throughput was recovered "
         "exactly; its training-quality curve is shown on a fair tokens axis. The one thing we "
         "cannot determine without a re-run is whether gpt2-xl fits on one GPU without "
         "checkpointing at the corrected batch size."),
        ("The micro-batch sweep does not hold the batch size fixed. ", "Because we increased "
         "micro-batches at a fixed micro-batch size, the total batch grew with M. So the speed "
         "rise mixes the shrinking bubble with a larger batch. The fixed-M=8 comparisons are "
         "unaffected."),
        ("PyTorch pipeline has no perplexity evaluation. ", "We compare its training-loss curve "
         "instead; perplexity is reported for the other two systems."),
        ("The output head is 'untied' in the pipeline. ", "A small technical necessity that gives "
         "the pipeline model ~5% more parameters than the tied single-GPU model — a minor "
         "difference we disclose."),
        ("Short runs and only two GPUs. ", "Speed/memory use ~30 steps; conclusions are specific "
         "to two stages over PCIe and should not be extrapolated to large clusters."),
    ])

    # ====================================================================== #
    heading(doc, "10. Conclusion", level=1)
    body(doc, "On two PCIe-connected T4 GPUs, pipeline parallelism does what the memory model "
              "predicts: it enables training a model (opt-2.7b) that a single 16 GB GPU cannot "
              "hold, at ~10.6 GB per GPU. The pipeline bubble follows theory, shrinking as we add "
              "micro-batches, with a small measurable overhead from the slow interconnect. Plain "
              "1F1B is the sweet spot — same bubble as GPipe but flat memory — while the newest "
              "zero-bubble schedules add only ~6% here because PCIe cannot hide their extra work. "
              "Between frameworks, PyTorch's pipelining is faster and leaner (1,835 vs 1,656 "
              "tokens/s; 8.4 vs 11.5 GB) while DeepSpeed offers more automation. Against a "
              "single-GPU gradient-checkpointing baseline (with its measurement corrected), the "
              "two-GPU pipeline is ~1.5× faster at lower per-GPU memory. The overarching lesson: at "
              "this scale and on this hardware, pipeline parallelism is primarily a way to fit "
              "bigger models — a memory win and a modest speed win — rather than a route to large "
              "speed-ups. All four assignment deliverables are addressed.")

    # ====================================================================== #
    heading(doc, "11. How to Reproduce", level=1)
    body(doc, "From the project root, on a 2× T4 machine (e.g. a Kaggle notebook):")
    code = doc.add_paragraph()
    cr = code.add_run(
        "# 0. prove the model split is correct\n"
        "python src/check_split.py --model gpt2-xl --num-stages 2\n\n"
        "# 1. run the full benchmark sweep\n"
        "bash bench/run_kaggle.sh\n\n"
        "# 2. build the tables and figures\n"
        "python bench/aggregate_results.py --model gpt2-xl --seq-len 512\n"
        "python bench/plots.py            --model gpt2-xl --seq-len 512\n"
        "python bench/plot_analysis.py    --model gpt2-xl --seq-len 512\n\n"
        "# 3. regenerate this Word report\n"
        "uv run --no-project python bench/make_docx.py")
    cr.font.name = "Consolas"; cr.font.size = Pt(9.5)

    # ====================================================================== #
    heading(doc, "12. References", level=1)
    refs = [
        "Huang et al. GPipe: Efficient Training of Giant Neural Networks using Pipeline Parallelism. NeurIPS 2019.",
        "Narayanan et al. PipeDream: Generalized Pipeline Parallelism for DNN Training. SOSP 2019.",
        "Narayanan et al. Efficient Large-Scale Language Model Training on GPU Clusters Using Megatron-LM. SC 2021.",
        "Qi et al. Zero Bubble Pipeline Parallelism. ICLR 2024.",
        "DeepSeek-AI. DeepSeek-V3 Technical Report. 2024/2025 (DualPipe).",
        "Rajbhandari et al. ZeRO: Memory Optimizations Toward Training Trillion Parameter Models. SC 2020.",
        "Chen et al. Training Deep Nets with Sublinear Memory Cost. 2016 (gradient checkpointing).",
        "Dettmers et al. 8-bit Optimizers via Block-wise Quantization. ICLR 2022.",
        "Radford et al. Language Models are Unsupervised Multitask Learners. 2019 (GPT-2).",
        "Zhang et al. OPT: Open Pre-trained Transformer Language Models. 2022.",
        "Merity et al. Pointer Sentinel Mixture Models. 2016 (WikiText).",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.add_run(f"[{i}] ").bold = True
        p.add_run(r).font.size = Pt(10)

    # ====================================================================== #
    heading(doc, "Appendix A — Environment", level=1)
    bullets(doc, [
        "GPUs: 2× NVIDIA T4 (16 GB), Kaggle, PCIe / no NVLink, Turing architecture.",
        "PyTorch: 2.10.0+cu128.",
        "Dataset: Salesforce/wikitext (wikitext-2-raw-v1), sequence length 512.",
        "Optimizer: 8-bit AdamW; precision fp16; random seed 1234.",
        "To complete: exact deepspeed / transformers / bitsandbytes versions from the Kaggle "
        "session (run: pip freeze | grep -E 'deepspeed|transformers|bitsandbytes').",
    ])

    heading(doc, "Appendix B — Full Results Table", level=1)
    body(doc, "Every run logged to results/results.csv (one row per experiment):")
    _full_results_table(doc)

    doc.save(OUT)
    print("wrote", OUT)


def _full_results_table(doc):
    if not os.path.exists(RESULTS_CSV):
        body(doc, "[results.csv not found]", italic=True)
        return
    with open(RESULTS_CSV, newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    cols = [("system", "system"), ("schedule", "sched"), ("num_micro_batches", "M"),
            ("global_batch", "batch"), ("tokens_per_sec", "tok/s"),
            ("peak_mem_gb_max", "GB/GPU"), ("mfu", "MFU"), ("final_loss", "loss"),
            ("eval_ppl", "ppl")]
    headers = [c[1] for c in cols]
    data = [[_fmt(r.get(k, ""), k) for k, _ in cols] for r in rows]
    table(doc, headers, data)


def _fmt(v, key):
    try:
        fv = float(v)
    except (ValueError, TypeError):
        return v
    if key == "tokens_per_sec":
        return f"{fv:,.0f}"
    if key in ("peak_mem_gb_max", "final_loss", "eval_ppl"):
        return f"{fv:.2f}"
    if key == "mfu":
        return f"{fv*100:.1f}%"
    if key in ("num_micro_batches", "global_batch"):
        return f"{int(fv)}"
    return f"{fv:g}"


if __name__ == "__main__":
    build()
